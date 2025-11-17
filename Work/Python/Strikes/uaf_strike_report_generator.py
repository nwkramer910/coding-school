#!/usr/bin/env python3
"""
Ukrainian Air Force Strike Report Generator

Extracts data from Ukrainian Air Force Telegram posts, populates an Excel tracking
sheet, and generates formatted Word document analysis.
"""

import re
import os
from datetime import datetime
from typing import Dict, Optional, Tuple, List
from openpyxl import Workbook, load_workbook
from openpyxl.styles import PatternFill
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX


class StrikeDataParser:
    """Parses UAF Telegram post text to extract strike data."""

    def __init__(self, text: str):
        self.text = text
        self.data = {}

    def parse_all(self) -> Dict:
        """Parse all data fields from the text."""
        self.data['date_phrase'] = self._parse_date_phrase()
        self.data['total_drones'] = self._parse_total_drones()
        self.data['shahed_count'] = self._parse_shahed_count()
        self.data['drones_shot_down'] = self._parse_drones_shot_down()

        # Parse missile data
        self.data['iskanders_fired'] = self._parse_iskanders_fired()
        self.data['iskanders_shot_down'] = self._parse_iskanders_shot_down()
        self.data['s300_400_fired'] = self._parse_missile_count(r'S-[34]00', 'fired')
        self.data['s300_400_shot_down'] = self._parse_missile_count(r'S-[34]00', 'shot down')
        self.data['kh101_fired'] = self._parse_missile_count(r'Kh-10[15]|Kh-555|Kh-55', 'fired')
        self.data['kh101_shot_down'] = self._parse_missile_count(r'Kh-10[15]|Kh-555|Kh-55', 'shot down')
        self.data['kh59_fired'] = self._parse_missile_count(r'Kh-[56]9', 'fired')
        self.data['kh59_shot_down'] = self._parse_missile_count(r'Kh-[56]9', 'shot down')
        self.data['kh31_fired'] = self._parse_missile_count(r'Kh-3[12]|Kh-22', 'fired')
        self.data['kh31_shot_down'] = self._parse_missile_count(r'Kh-3[12]|Kh-22', 'shot down')
        self.data['kalibr_fired'] = self._parse_missile_count(r'["\']?Kalibr["\']?', 'fired')
        self.data['kalibr_shot_down'] = self._parse_missile_count(r'["\']?Kalibr["\']?', 'shot down')
        self.data['kinzhal_fired'] = self._parse_missile_count(r'["\']?Kinzhal["\']?|Kh-47M?2?', 'fired')
        self.data['kinzhal_shot_down'] = self._parse_missile_count(r'["\']?Kinzhal["\']?|Kh-47M?2?', 'shot down')

        # Calculate totals
        self.data['total_missiles'] = self._calculate_total_missiles()
        self.data['total_air_targets'] = self._calculate_total_air_targets()

        # Parse weapon-specific locations
        self.data['iskander_m_locations'] = self._parse_weapon_locations(r'Iskander-M/KN-23')
        self.data['iskander_k_locations'] = self._parse_weapon_locations(r'Iskander-K')
        self.data['kinzhal_locations'] = self._parse_weapon_locations(r'Kh-47M?2?\s*["\']?Kinzhal["\']?|["\']?Kinzhal["\']?')
        self.data['kalibr_locations'] = self._parse_weapon_locations(r'["\']?Kalibr["\']?')
        self.data['kh101_locations'] = self._parse_weapon_locations(r'Kh-10[15]|Kh-555|Kh-55')
        self.data['kh59_locations'] = self._parse_weapon_locations(r'Kh-[56]9')
        self.data['kh31_locations'] = self._parse_weapon_locations(r'Kh-3[12]|Kh-22')
        self.data['s300_400_locations'] = self._parse_weapon_locations(r'S-[34]00')
        self.data['drone_launch_locations'] = self._parse_drone_locations()

        # Parse impact info
        self.data['impact_info'] = self._parse_impact_info()

        return self.data

    def _parse_date_phrase(self) -> Optional[str]:
        """Extract date phrase like 'October 21 and 22'."""
        # Look for patterns like "night of October 21-22" or "October 21 and 22"
        patterns = [
            r'night of (\w+ \d+(?:\s*(?:and|-)\s*\d+)?)',
            r'on (\w+ \d+(?:\s*(?:and|-)\s*\d+)?)',
            r'(\w+ \d+(?:\s*(?:and|-)\s*\d+)?)'
        ]

        for pattern in patterns:
            match = re.search(pattern, self.text, re.IGNORECASE)
            if match:
                date_str = match.group(1)
                # Normalize to "Month DD and DD" format
                date_str = date_str.replace('-', ' and ')
                return date_str
        return None

    def _parse_total_drones(self) -> Optional[int]:
        """Extract total number of drones launched."""
        # Look for patterns in the summary line or bullet points
        patterns = [
            # "503 air attack weapons – 45 missiles and 458 UAVs"
            r'(\d+)\s+missiles.*?and\s+(\d+)\s+UAVs',
            # "- 458 Shahed and Gerbera strike UAVs"
            r'-\s*(\d+)\s+(?:Shahed|strike\s+UAVs)',
            # Legacy patterns
            r'launched\s+(\d+)\s+(?:Shahed|drones|UAVs)',
            r'(\d+)\s+Shahed-type.*?drones',
        ]

        for pattern in patterns:
            match = re.search(pattern, self.text, re.IGNORECASE)
            if match:
                # For the first pattern, we want the UAV count (second group)
                if 'missiles.*?and' in pattern:
                    return int(match.group(2))
                else:
                    return int(match.group(1))
        return None

    def _parse_shahed_count(self) -> Optional[int]:
        """Extract count of Shahed drones specifically."""
        # Look for "about X of them were Shahed UAVs" or similar
        patterns = [
            r'about\s+(\d+)\s+of\s+them\s+were\s+Shahed',
            r'roughly\s+(\d+)\s+of\s+(?:which\s+were|these\s+were|them\s+were)\s+Shahed',
            r'(\d+)\s+of\s+(?:which|them)\s+were\s+Shahed',
        ]

        for pattern in patterns:
            match = re.search(pattern, self.text, re.IGNORECASE)
            if match:
                return int(match.group(1))

        # If no specific breakdown found, return None (not total - that would be incorrect)
        return None

    def _parse_drones_shot_down(self) -> Optional[int]:
        """Extract number of drones shot down."""
        patterns = [
            # "- 406 Shahed and Gerbera strike UAVs"
            r'shot down[^:]*:\s*-\s*(\d+)\s+(?:Shahed|strike\s+UAVs)',
            # Legacy patterns
            r'downed\s+(\d+)\s+(?:drones|UAVs)',
            r'shot\s+down\s+(\d+)\s+(?:drones|UAVs)',
            r'(\d+)\s+(?:drones|UAVs).*?(?:downed|destroyed)'
        ]

        for pattern in patterns:
            match = re.search(pattern, self.text, re.IGNORECASE | re.DOTALL)
            if match:
                return int(match.group(1))
        return None

    def _parse_iskanders_fired(self) -> Optional[str]:
        """Parse Iskander missiles in format 'X Isk-K, Y Isk-M/KN-23'."""
        isk_k_count = self._parse_missile_count(r'Iskander-K|Isk-K', 'fired')
        isk_m_count = self._parse_missile_count(r'Iskander-M|Isk-M|KN-23', 'fired')

        if isk_k_count is None and isk_m_count is None:
            return None

        parts = []
        if isk_k_count:
            parts.append(f"{isk_k_count} Isk-K")
        if isk_m_count:
            parts.append(f"{isk_m_count} Isk-M/KN-23")

        return ", ".join(parts) if parts else None

    def _parse_iskanders_shot_down(self) -> Optional[str]:
        """Parse Iskander shootdowns in format 'X Isk-K, Y Isk-M/KN-23'."""
        isk_k_count = self._parse_missile_count(r'Iskander-K|Isk-K', 'shot down')
        isk_m_count = self._parse_missile_count(r'Iskander-M|Isk-M|KN-23', 'shot down')

        if isk_k_count is None and isk_m_count is None:
            return None

        parts = []
        if isk_k_count:
            parts.append(f"{isk_k_count} Isk-K")
        if isk_m_count:
            parts.append(f"{isk_m_count} Isk-M/KN-23")

        return ", ".join(parts) if parts else None

    def _parse_missile_count(self, missile_pattern: str, action: str) -> Optional[int]:
        """Generic missile count parser for fired/shot down."""
        if action == 'fired':
            # Look for bullet point format: "- X missile-type"
            # Split by "shot down" or "repelled" to only look in the launch section
            split_keywords = ['shot down', 'repelled', 'air defense']
            text_section = self.text
            for keyword in split_keywords:
                if keyword in self.text.lower():
                    text_section = self.text.split(keyword)[0]
                    break

            # Try bullet point format first: "- 25 Iskander-M/KN-23"
            bullet_pattern = rf'-\s*(\d+)\s+(?:{missile_pattern})'
            matches = re.findall(bullet_pattern, text_section, re.IGNORECASE)
            if matches:
                return sum(int(m) for m in matches)

            # Try inline format
            inline_pattern = rf'(\d+)\s+(?:{missile_pattern})'
            matches = re.findall(inline_pattern, text_section, re.IGNORECASE)
            if matches:
                return sum(int(m) for m in matches)

        else:  # shot down
            # Note: UAF posts often don't break down missile shootdowns by type
            # They may just say "9 missiles of various types"
            # Look in the shootdown section after "shot down" keywords
            if 'shot down' in self.text.lower():
                text_section = self.text.split('shot down')[1]
            else:
                text_section = self.text

            # Try to find specific missile type in shootdown section
            patterns = [
                rf'-\s*(\d+)\s+(?:{missile_pattern})',
                rf'(\d+)\s+(?:{missile_pattern})',
            ]

            for pattern in patterns:
                matches = re.findall(pattern, text_section, re.IGNORECASE)
                if matches:
                    return sum(int(m) for m in matches)

        return None

    def _calculate_total_missiles(self) -> Optional[int]:
        """Calculate total missiles from all types."""
        # First try to parse from summary line: "503 air attack weapons – 45 missiles"
        summary_pattern = r'(\d+)\s+missiles.*?and\s+\d+\s+UAVs'
        match = re.search(summary_pattern, self.text, re.IGNORECASE)
        if match:
            return int(match.group(1))

        # Otherwise calculate from individual types
        missile_counts = [
            self._extract_iskander_total(self.data.get('iskanders_fired')),
            self.data.get('s300_400_fired', 0) or 0,
            self.data.get('kh101_fired', 0) or 0,
            self.data.get('kh59_fired', 0) or 0,
            self.data.get('kh31_fired', 0) or 0,
            self.data.get('kalibr_fired', 0) or 0,
            self.data.get('kinzhal_fired', 0) or 0,
        ]

        total = sum(missile_counts)
        return total if total > 0 else None

    def _calculate_total_air_targets(self) -> Optional[int]:
        """Calculate total air targets (drones + missiles)."""
        # First try to parse from summary line: "503 air attack weapons"
        summary_pattern = r'(\d+)\s+air\s+attack\s+weapons'
        match = re.search(summary_pattern, self.text, re.IGNORECASE)
        if match:
            return int(match.group(1))

        # Otherwise calculate from drones + missiles
        drones = self.data.get('total_drones', 0) or 0
        missiles = self.data.get('total_missiles', 0) or 0
        total = drones + missiles
        return total if total > 0 else None

    def _extract_iskander_total(self, iskander_str: Optional[str]) -> int:
        """Extract total count from Iskander string format."""
        if not iskander_str:
            return 0
        numbers = re.findall(r'(\d+)', iskander_str)
        return sum(int(n) for n in numbers)

    def _parse_weapon_locations(self, weapon_pattern: str) -> Optional[str]:
        """Parse launch locations for a specific weapon type from parentheses.

        Example: "- 25 Iskander-M/KN-23 ballistic missiles (from Kursk, Voronezh, and Rostov regions – Russian Federation);"
        Returns: "Kursk, Voronezh, and Rostov oblasts"
        """
        # Find the line containing this weapon type
        pattern = rf'-\s*\d+\s+(?:{weapon_pattern}).*?\(from\s+([^)]+)\)'
        match = re.search(pattern, self.text, re.IGNORECASE | re.DOTALL)

        if match:
            location_text = match.group(1)
            # Clean up: remove "– Russian Federation", "– RF", etc.
            location_text = re.sub(r'\s*[–-]\s*(?:Russian Federation|RF|TOT AR Crimea).*$', '', location_text)
            # Convert "regions" to "oblasts", "region" to "Oblast"
            location_text = re.sub(r'\bregions?\b', lambda m: 'oblasts' if m.group().endswith('s') else 'Oblast', location_text, flags=re.IGNORECASE)
            # Clean up extra whitespace
            location_text = re.sub(r'\s+', ' ', location_text).strip()
            return location_text

        return None

    def _parse_drone_locations(self) -> Optional[str]:
        """Parse drone launch locations with special formatting.

        Example: "from the directions of Kursk, Millerovo, Orel, Primorsko-Akhtarsk – RF, Gvardeyskoye – TOT AR Crimea"
        Returns: "the directions of Kursk and Oryol cities; Millerovo, Rostov Oblast; Primorsko-Akhtarsk, Krasnodar Krai; and occupied Hvardiiske, Crimea"

        Note: The analyst will need to manually format city vs oblast details as the automated translation doesn't preserve this.
        """
        # Find the drone line with locations
        pattern = r'-\s*\d+\s+Shahed.*?\((?:other types of drones\))?\s*from\s+([^)]+)\)'
        match = re.search(pattern, self.text, re.IGNORECASE | re.DOTALL)

        if match:
            location_text = match.group(1)
            # Clean up: remove "– Russian Federation", "– RF", etc.
            location_text = re.sub(r'\s*[–-]\s*(?:Russian Federation|RF)(?:\s*,)?', '', location_text)
            location_text = re.sub(r'\s*[–-]\s*TOT AR Crimea', '', location_text)
            location_text = re.sub(r'\s+', ' ', location_text).strip()
            # Preserve "the directions of" if present
            if 'the directions of' not in location_text.lower() and 'directions of' in location_text.lower():
                location_text = re.sub(r'directions of\s+', 'the directions of ', location_text, flags=re.IGNORECASE)
            return location_text

        return None

    def _parse_launch_locations(self, weapon_type: str) -> Optional[str]:
        """Parse launch locations for drones or missiles (legacy method)."""
        # Look for "from [location]" patterns
        patterns = [
            r'from\s+((?:[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*(?:,\s*)?)+)',
            r'launched\s+from\s+((?:[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*(?:,\s*)?)+)',
        ]

        locations = []
        for pattern in patterns:
            matches = re.findall(pattern, self.text)
            locations.extend(matches)

        return ", ".join(set(locations)) if locations else None

    def _parse_impact_info(self) -> Dict[str, Optional[str]]:
        """Parse impact, damage, and target information."""
        info = {
            'hits': None,
            'debris': None,
            'lost_drones': None,
            'target_oblasts': None,
            'damage': None
        }

        # Parse hits - new format: "26 missiles and 52 strike UAVs have been recorded as hitting 25 locations"
        hit_patterns = [
            r'(\d+)\s+missiles?\s+and\s+(\d+)\s+(?:strike\s+)?(?:UAVs|drones).*?hitting\s+(\d+)\s+locations',
            r'(\d+)\s+missiles?\s+and\s+(\d+)\s+(?:strike\s+)?(?:UAVs|drones)\s+hit\s+(.+?)(?:\.|,)',
        ]
        for pattern in hit_patterns:
            hit_match = re.search(pattern, self.text, re.IGNORECASE)
            if hit_match:
                info['hits'] = f"{hit_match.group(1)} missiles and {hit_match.group(2)} drones hit {hit_match.group(3)}"
                break

        # Parse debris - new format: "debris from downed missiles has been recorded at 4 locations"
        debris_patterns = [
            r'debris.*?recorded\s+at\s+(\d+)\s+locations',
            r'debris.*?fell\s+(?:on|in|at)\s+(.+?)(?:\.|,)',
        ]
        for pattern in debris_patterns:
            debris_match = re.search(pattern, self.text, re.IGNORECASE)
            if debris_match:
                info['debris'] = debris_match.group(1)
                break

        # Parse lost drones
        lost_match = re.search(r'(\d+)\s+(?:UAVs|drones)?\s+(?:did\s+not\s+reach|lost|were\s+lost)', self.text, re.IGNORECASE)
        if lost_match:
            info['lost_drones'] = lost_match.group(1)

        # Parse target oblasts - new format: "The main targets of the attack were Kyiv, Dnipropetrovsk, and Poltava regions"
        oblast_patterns = [
            r'main\s+targets.*?were\s+(.+?)\s+regions?',
            r'targeted\s+(.+?)(?:\.|and\s+also\s+affected)',
        ]
        for pattern in oblast_patterns:
            oblast_match = re.search(pattern, self.text, re.IGNORECASE)
            if oblast_match:
                info['target_oblasts'] = oblast_match.group(1)
                break

        # Parse damage
        damage_match = re.search(r'damaged\s+(.+?)\s+infrastructure\s+in\s+(.+?)(?:\.|,)', self.text, re.IGNORECASE)
        if damage_match:
            info['damage'] = f"{damage_match.group(1)} infrastructure in {damage_match.group(2)}"

        return info


class NumberFormatter:
    """Formats numbers according to style guide."""

    SPELLED_NUMBERS = {
        1: 'one', 2: 'two', 3: 'three', 4: 'four', 5: 'five',
        6: 'six', 7: 'seven', 8: 'eight', 9: 'nine'
    }

    @classmethod
    def format(cls, num: Optional[int]) -> str:
        """Format number: spell out 1-9, use numerals for 10+, NULL for None."""
        if num is None:
            return '**NULL**'
        if num <= 9:
            return cls.SPELLED_NUMBERS.get(num, str(num))
        return str(num)


class ExcelHandler:
    """Handles Excel file operations."""

    HEADERS = [
        'Date', 'Total Drones', 'Total Missiles', 'Total Air Targets',
        'Shahed Count', 'Drones Shot Down', 'Iskanders Fired', 'Iskanders Shot Down',
        'S-300/400 Fired', 'S-300/400 Shot Down', 'Kh-101/55/555 Fired', 'Kh-101/55/555 Shot Down',
        'Kh-59/69 Fired', 'Kh-59/69 Shot Down', 'Kh-31/32/22 Fired', 'Kh-31/32/22 Shot Down',
        'Kalibr Fired', 'Kalibr Shot Down', 'Kh-47 Kinzhal Fired', 'Kh-47 Kinzhal Shot Down',
        'Misc. Munitions', '% Intercepted Missiles', 'Link'
    ]

    def __init__(self, filepath: str = 'RUAF_Strikes_in_UA.xlsx'):
        self.filepath = filepath
        self.workbook = None
        self.worksheet = None

    def load_or_create(self):
        """Load existing workbook or create new one."""
        if os.path.exists(self.filepath):
            self.workbook = load_workbook(self.filepath)
            self.worksheet = self.workbook.active
        else:
            self.workbook = Workbook()
            self.worksheet = self.workbook.active
            self.worksheet.append(self.HEADERS)

    def append_data(self, data: Dict):
        """Append new row of data to the worksheet."""
        row_data = [
            data.get('date_phrase'),
            data.get('total_drones'),
            data.get('total_missiles'),
            data.get('total_air_targets'),
            data.get('shahed_count'),
            data.get('drones_shot_down'),
            data.get('iskanders_fired'),
            data.get('iskanders_shot_down'),
            data.get('s300_400_fired'),
            data.get('s300_400_shot_down'),
            data.get('kh101_fired'),
            data.get('kh101_shot_down'),
            data.get('kh59_fired'),
            data.get('kh59_shot_down'),
            data.get('kh31_fired'),
            data.get('kh31_shot_down'),
            data.get('kalibr_fired'),
            data.get('kalibr_shot_down'),
            data.get('kinzhal_fired'),
            data.get('kinzhal_shot_down'),
            '',  # Misc. Munitions (manual)
            self._calculate_interception_rate(data),
            ''  # Link (manual)
        ]

        self.worksheet.append(row_data)

    def _calculate_interception_rate(self, data: Dict) -> Optional[float]:
        """Calculate missile interception percentage."""
        total_missiles = data.get('total_missiles')
        if not total_missiles or total_missiles == 0:
            return None

        # Sum all shot down missiles (including parsed Iskanders)
        shot_down = sum([
            data.get('s300_400_shot_down', 0) or 0,
            data.get('kh101_shot_down', 0) or 0,
            data.get('kh59_shot_down', 0) or 0,
            data.get('kh31_shot_down', 0) or 0,
            data.get('kalibr_shot_down', 0) or 0,
            data.get('kinzhal_shot_down', 0) or 0,
            self._extract_iskander_total(data.get('iskanders_shot_down')),
        ])

        return round((shot_down / total_missiles) * 100, 2)

    def _extract_iskander_total(self, iskander_str: Optional[str]) -> int:
        """Extract total count from Iskander string format."""
        if not iskander_str:
            return 0
        numbers = re.findall(r'(\d+)', iskander_str)
        return sum(int(n) for n in numbers)

    def save(self):
        """Save the workbook."""
        self.workbook.save(self.filepath)
        print(f"Excel file saved: {self.filepath}")


class WordDocumentGenerator:
    """Generates formatted Word document report."""

    def __init__(self, data: Dict, original_text: str = ''):
        self.data = data
        self.text = original_text
        self.doc = Document()

    def generate(self) -> str:
        """Generate the Word document and return filename."""
        # Generate filename from date
        filename = self._generate_filename()

        # Add content in published ISW format
        self._add_missiles_and_drones_sentence()
        self._add_drone_details_sentence()
        self._add_combined_results_sentence()

        # Save document
        self.doc.save(filename)
        print(f"Word document saved: {filename}")
        return filename

    def _generate_filename(self) -> str:
        """Generate filename like Strike_Oct22_2024.docx."""
        date_phrase = self.data.get('date_phrase')
        if date_phrase:
            # Parse date from phrase
            match = re.search(r'(\w+)\s+(\d+)', date_phrase)
            if match:
                month = match.group(1)[:3]  # First 3 letters
                day = match.group(2)
                year = datetime.now().year
                return f"Strike_{month}{day}_{year}.docx"

        # Fallback to current date
        now = datetime.now()
        return f"Strike_{now.strftime('%b%d_%Y')}.docx"

    def _add_missiles_and_drones_sentence(self):
        """First sentence: 'The Ukrainian Air Force reported that Russian forces launched X drones and missiles, including [missiles].'"""
        total = NumberFormatter.format(self.data.get('total_air_targets'))

        # Build missile list
        missiles = self._build_missile_list_with_locations()

        if missiles:
            missile_list = self._format_list_with_semicolons(missiles)
            text = f"The Ukrainian Air Force reported that Russian forces launched {total} drones and missiles, including {missile_list}."
        else:
            text = f"The Ukrainian Air Force reported that Russian forces launched {total} drones."

        self._add_paragraph(text)

    def _add_drone_details_sentence(self):
        """Second sentence: 'The Ukrainian Air Force reported that Russian forces also launched X Shahed-type...'"""
        total_drones = NumberFormatter.format(self.data.get('total_drones'))
        shahed_count = NumberFormatter.format(self.data.get('shahed_count'))
        locations = self._format_with_null(self.data.get('drone_launch_locations'))

        if shahed_count and shahed_count != '**NULL**':
            text = (f"The Ukrainian Air Force reported that Russian forces also launched {total_drones} "
                   f"Shahed-type, Gerbera-type, and other drones – of which roughly {shahed_count} were "
                   f"Shahed-type drones – from {locations}.")
        else:
            text = (f"The Ukrainian Air Force reported that Russian forces also launched {total_drones} "
                   f"Shahed-type, Gerbera-type, and other drones from {locations}.")

        self._add_paragraph(text)

    def _add_combined_results_sentence(self):
        """Third sentence: Combined shootdowns, hits, and debris in one sentence."""
        drones_down = NumberFormatter.format(self.data.get('drones_shot_down'))
        missiles_down = NumberFormatter.format(self.data.get('total_missiles'))  # This needs to be parsed from "9 missiles"

        # Parse actual missiles shot down count
        impact_info = self.data.get('impact_info', {})

        # Build the combined sentence parts
        parts = []

        # Part 1: Shootdowns
        # For now, check if we have generic "X missiles of various types" or specific countsp
        if missiles_down and missiles_down != '**NULL**':
            # UAF reports "9 missiles of various types" - extract this
            missile_down_text = self._parse_total_missiles_shot_down()
            if missile_down_text:
                parts.append(f"downed {drones_down} drones and {missile_down_text} missiles")
            else:
                parts.append(f"downed {drones_down} drones")
        else:
            parts.append(f"downed {drones_down} drones")

        # Part 2: Hits
        hits_info = impact_info.get('hits')
        if hits_info:
            parts.append(hits_info.replace(" hit ", " struck "))

        # Part 3: Debris
        debris_info = impact_info.get('debris')
        if debris_info:
            debris_count = NumberFormatter.format(int(debris_info)) if debris_info.isdigit() else debris_info
            parts.append(f"downed debris fell on {debris_count} locations")

        # Combine all parts
        if parts:
            combined = ', '.join(parts[:-1]) + f", and {parts[-1]}" if len(parts) > 1 else parts[0]
            text = f"The Ukrainian Air Force reported that Ukrainian air defense {combined}."
        else:
            text = "The Ukrainian Air Force reported strike results."

        self._add_paragraph(text)

    def _parse_total_missiles_shot_down(self) -> Optional[str]:
        """Parse total missiles shot down from text (e.g., 'nine' from '9 missiles of various types')."""
        # Look for "- 9 missiles of various types" in the shootdown section
        pattern = r'shot down[^:]*:\s*.*?-\s*(\d+)\s+missiles?\s+of\s+various\s+types'
        match = re.search(pattern, self.text if hasattr(self, 'text') else '', re.IGNORECASE | re.DOTALL)
        if match:
            count = int(match.group(1))
            return NumberFormatter.format(count)
        return None

    def _build_missile_list_with_locations(self) -> List[str]:
        """Build list of missiles with their specific locations for sentence 1."""
        missiles = []

        # Iskander-M/KN-23
        iskanders_fired = self.data.get('iskanders_fired')
        if iskanders_fired:
            isk_m_match = re.search(r'(\d+)\s+Isk-M/KN-23', iskanders_fired)
            if isk_m_match:
                count = NumberFormatter.format(int(isk_m_match.group(1)))
                locations = self._format_with_null(self.data.get('iskander_m_locations'))
                missiles.append(f"{count} Iskander-M/KN-23 ballistic missiles from {locations}")

        # Iskander-K
            isk_k_match = re.search(r'(\d+)\s+Isk-K', iskanders_fired)
            if isk_k_match:
                count = NumberFormatter.format(int(isk_k_match.group(1)))
                locations = self._format_with_null(self.data.get('iskander_k_locations'))
                missiles.append(f"{count} Iskander-K cruise missiles from {locations}")

        # Kh-47M2 Kinzhal
        if self.data.get('kinzhal_fired'):
            count = NumberFormatter.format(self.data['kinzhal_fired'])
            locations = self._format_with_null(self.data.get('kinzhal_locations'))
            missiles.append(f"{count} Kh-47M2 Kinzhal aeroballistic missiles from {locations}")

        # Kalibr
        if self.data.get('kalibr_fired'):
            count = NumberFormatter.format(self.data['kalibr_fired'])
            locations = self._format_with_null(self.data.get('kalibr_locations'))
            missiles.append(f"{count} Kalibr missiles from {locations}")

        # Kh-101/55/555
        if self.data.get('kh101_fired'):
            count = NumberFormatter.format(self.data['kh101_fired'])
            locations = self._format_with_null(self.data.get('kh101_locations'))
            missiles.append(f"{count} Kh-101/Kh-555/Kh-55 cruise missiles from {locations}")

        # Kh-59/69
        if self.data.get('kh59_fired'):
            count = NumberFormatter.format(self.data['kh59_fired'])
            locations = self._format_with_null(self.data.get('kh59_locations'))
            missiles.append(f"{count} Kh-59/Kh-69 guided air missiles from {locations}")

        # Kh-31/32/22
        if self.data.get('kh31_fired'):
            count = NumberFormatter.format(self.data['kh31_fired'])
            locations = self._format_with_null(self.data.get('kh31_locations'))
            missiles.append(f"{count} Kh-31P/Kh-32/Kh-22 cruise missiles from {locations}")

        # S-300/400
        if self.data.get('s300_400_fired'):
            count = NumberFormatter.format(self.data['s300_400_fired'])
            locations = self._format_with_null(self.data.get('s300_400_locations'))
            missiles.append(f"{count} S-300/S-400 SAMs from {locations}")

        return missiles

    def _format_with_null(self, value: Optional[str]) -> str:
        """Format value, returning **NULL** if None."""
        return value if value is not None else '**NULL**'

    def _format_list_with_semicolons(self, items: List[str]) -> str:
        """Format list with semicolons and 'and' before last item."""
        if not items:
            return ''
        if len(items) == 1:
            return items[0]
        if len(items) == 2:
            return f"{items[0]} and {items[1]}"
        return '; '.join(items[:-1]) + f"; and {items[-1]}"

    def _add_paragraph(self, text: str):
        """Add paragraph with NULL highlighting."""
        para = self.doc.add_paragraph()

        # Split text by **NULL** markers
        parts = text.split('**NULL**')

        for i, part in enumerate(parts):
            # Add regular text
            run = para.add_run(part)

            # Add highlighted NULL if not last part
            if i < len(parts) - 1:
                null_run = para.add_run('NULL')
                null_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                null_run.font.bold = True


def main():
    """Main execution function."""
    print("=== UAF Strike Report Generator ===\n")
    print("Paste the English-translated UAF Telegram post text below.")
    print("When finished, press Ctrl+D (Linux/Mac) or Ctrl+Z then Enter (Windows):\n")

    # Read multiline input
    lines = []
    try:
        while True:
            line = input()
            lines.append(line)
    except EOFError:
        pass

    text = '\n'.join(lines)

    if not text.strip():
        print("Error: No text provided.")
        return

    print("\n--- Parsing data... ---")
    parser = StrikeDataParser(text)
    data = parser.parse_all()

    print("\n--- Updating Excel file... ---")
    excel_handler = ExcelHandler()
    excel_handler.load_or_create()
    excel_handler.append_data(data)
    excel_handler.save()

    print("\n--- Generating Word document... ---")
    doc_generator = WordDocumentGenerator(data, text)
    filename = doc_generator.generate()

    print("\n=== Processing Complete ===")
    print(f"Excel file: RUAF_Strikes_in_UA.xlsx")
    print(f"Word document: {filename}")


if __name__ == '__main__':
    main()
