#!/usr/bin/env python3
"""
Fighting Circles Generator - GUI Version
Processes UA GEN STAFF sitreps and generates fighting circles with spatial matching.
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import json
import os
import re
from pathlib import Path
from datetime import datetime
import threading

# Import the core functionality from fcs.py
from fcs_base import (
    extract_gen_staff_settlements,
    match_settlements_to_buckets,
    create_fighting_circles,
    parse_document_structure
)
from docx import Document
import geopandas as gpd


class ConfigManager:
    """Manages persistent configuration settings."""

    def __init__(self, config_file=".fcs_config.json"):
        self.config_file = config_file
        self.config = self.load_config()

    def load_config(self):
        """Load configuration from file."""
        if os.path.exists(self.config_file):
            try:
                with open(self.config_file, 'r') as f:
                    return json.load(f)
            except:
                pass

        # Default configuration
        return {
            'shapefile_paths': {
                'city_shp': 'citypoints.shp',
                'frontline_shp': 'frontline_frontline.shp',
                'bucket_shp': 'collection_buckets.shp'
            },
            'bucket_field': 'Name',
            'bucket_to_document_map': {},
            'last_docx_dir': str(Path.home()),
            'distance_threshold': 15000
        }

    def save_config(self):
        """Save configuration to file."""
        try:
            with open(self.config_file, 'w') as f:
                json.dump(self.config, indent=2, fp=f)
            return True
        except Exception as e:
            print(f"Error saving config: {e}")
            return False

    def get(self, key, default=None):
        """Get configuration value."""
        keys = key.split('.')
        value = self.config
        for k in keys:
            if isinstance(value, dict):
                value = value.get(k, default)
            else:
                return default
        return value

    def set(self, key, value):
        """Set configuration value."""
        keys = key.split('.')
        config = self.config
        for k in keys[:-1]:
            if k not in config:
                config[k] = {}
            config = config[k]
        config[keys[-1]] = value
        self.save_config()


class HeadingMapperDialog:
    """Dialog for mapping shapefile buckets to document section paths."""

    def __init__(self, parent, bucket_names, current_mappings):
        self.result = None
        self.bucket_names = sorted(bucket_names)
        self.current_mappings = current_mappings.copy()

        # Create dialog window
        self.dialog = tk.Toplevel(parent)
        self.dialog.title("Map Shapefile Buckets to Document Sections")
        self.dialog.geometry("1000x600")
        self.dialog.transient(parent)
        self.dialog.grab_set()

        self.setup_ui()

        # Center on parent
        self.dialog.update_idletasks()
        x = parent.winfo_x() + (parent.winfo_width() // 2) - (self.dialog.winfo_width() // 2)
        y = parent.winfo_y() + (parent.winfo_height() // 2) - (self.dialog.winfo_height() // 2)
        self.dialog.geometry(f"+{x}+{y}")

    def setup_ui(self):
        """Setup the dialog UI."""
        # Instructions
        instructions = tk.Label(
            self.dialog,
            text="For each shapefile bucket, specify the document section path to match:",
            font=('Arial', 10, 'bold'),
            pady=10
        )
        instructions.pack()

        # Sub-instructions
        sub_instructions = tk.Label(
            self.dialog,
            text='(Enter the full path from Heading 1 through bold bullets, e.g., "Kharkiv Oblast - Kupyansk direction")',
            font=('Arial', 9),
            pady=5
        )
        sub_instructions.pack()

        # Scrollable frame for mappings
        canvas = tk.Canvas(self.dialog)
        scrollbar = ttk.Scrollbar(self.dialog, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        # Create mapping widgets
        self.mapping_vars = {}

        for i, bucket_name in enumerate(self.bucket_names):
            frame = ttk.Frame(scrollable_frame, padding=5)
            frame.grid(row=i, column=0, sticky='ew', padx=10, pady=2)

            # Bucket name label (from shapefile)
            label = ttk.Label(frame, text=f"{bucket_name}:", width=30, anchor='w', font=('Arial', 9, 'bold'))
            label.grid(row=0, column=0, padx=(0, 10))

            # Document path entry (user enters the path)
            var = tk.StringVar(value=self.current_mappings.get(bucket_name, ''))
            entry = ttk.Entry(
                frame,
                textvariable=var,
                width=70
            )
            entry.grid(row=0, column=1, sticky='ew')

            self.mapping_vars[bucket_name] = var

            frame.columnconfigure(1, weight=1)

        canvas.pack(side="left", fill="both", expand=True, padx=10, pady=10)
        scrollbar.pack(side="right", fill="y")

        # Buttons
        button_frame = ttk.Frame(self.dialog)
        button_frame.pack(side='bottom', pady=10)

        ttk.Button(
            button_frame,
            text="Save Mappings",
            command=self.save_mappings
        ).pack(side='left', padx=5)

        ttk.Button(
            button_frame,
            text="Cancel",
            command=self.dialog.destroy
        ).pack(side='left', padx=5)

    def save_mappings(self):
        """Save the mappings and close dialog."""
        self.result = {
            bucket: var.get().strip()
            for bucket, var in self.mapping_vars.items()
            if var.get().strip()  # Only save non-empty mappings
        }
        self.dialog.destroy()

    def show(self):
        """Show dialog and return result."""
        self.dialog.wait_window()
        return self.result


class FightingCirclesGUI:
    """Main GUI application for Fighting Circles Generator."""

    def __init__(self, root):
        self.root = root
        self.root.title("Fighting Circles Generator")
        self.root.geometry("900x700")

        self.config = ConfigManager()
        self.current_docx_path = None
        self.processing = False

        self.setup_ui()

    def setup_ui(self):
        """Setup the main UI."""
        # Main container
        main_frame = ttk.Frame(self.root, padding=10)
        main_frame.pack(fill='both', expand=True)

        # Title
        title = tk.Label(
            main_frame,
            text="Fighting Circles Generator",
            font=('Arial', 16, 'bold')
        )
        title.pack(pady=(0, 10))

        # Document selection
        self.setup_document_section(main_frame)

        # Shapefile paths
        self.setup_shapefile_section(main_frame)

        # Mapping button
        self.setup_mapping_section(main_frame)

        # Run button
        self.setup_run_section(main_frame)

        # Output/Results area
        self.setup_output_section(main_frame)

    def setup_document_section(self, parent):
        """Setup document selection section."""
        frame = ttk.LabelFrame(parent, text="1. Select Document and Output", padding=10)
        frame.pack(fill='x', pady=5)

        self.docx_path_var = tk.StringVar()
        # Load last output directory from config if available
        last_output = self.config.get('last_output_dir', '')
        self.output_path_var = tk.StringVar(value=last_output)

        # Document path
        ttk.Label(frame, text="Daily Collect Document (.docx):").grid(
            row=0, column=0, sticky='w', pady=5
        )

        ttk.Entry(frame, textvariable=self.docx_path_var, width=60).grid(
            row=0, column=1, padx=5, sticky='ew'
        )

        ttk.Button(frame, text="Browse...", command=self.browse_document).grid(
            row=0, column=2, padx=5
        )

        # Output directory path
        ttk.Label(frame, text="Output Directory (optional):").grid(
            row=1, column=0, sticky='w', pady=5
        )

        ttk.Entry(frame, textvariable=self.output_path_var, width=60).grid(
            row=1, column=1, padx=5, sticky='ew'
        )

        ttk.Button(frame, text="Browse...", command=self.browse_output).grid(
            row=1, column=2, padx=5
        )

        # Info label
        ttk.Label(
            frame,
            text="(If not specified, output will be saved to the document's directory)",
            font=('Arial', 8),
            foreground='gray'
        ).grid(row=2, column=1, sticky='w', pady=(0, 5))

        frame.columnconfigure(1, weight=1)
        

    def setup_shapefile_section(self, parent):
        """Setup shapefile paths section."""
        frame = ttk.LabelFrame(parent, text="2. Shapefile Paths (Optional - defaults to script directory)", padding=10)
        frame.pack(fill='x', pady=5)

        self.shapefile_vars = {}

        shapefiles = [
            ('city_shp', 'City Points Shapefile:'),
            ('frontline_shp', 'Frontline Shapefile:'),
            ('bucket_shp', 'Collection Buckets Shapefile:')
        ]

        for i, (key, label) in enumerate(shapefiles):
            ttk.Label(frame, text=label).grid(row=i, column=0, sticky='w', pady=2)

            var = tk.StringVar(value=self.config.get(f'shapefile_paths.{key}', ''))
            self.shapefile_vars[key] = var

            ttk.Entry(frame, textvariable=var, width=50).grid(
                row=i, column=1, padx=5, sticky='ew'
            )

            ttk.Button(
                frame,
                text="Browse...",
                command=lambda k=key: self.browse_shapefile(k)
            ).grid(row=i, column=2, padx=5)

        frame.columnconfigure(1, weight=1)

    def setup_mapping_section(self, parent):
        """Setup bucket to document mapping section."""
        frame = ttk.LabelFrame(parent, text="3. Bucket to Document Mapping", padding=10)
        frame.pack(fill='x', pady=5)

        ttk.Label(
            frame,
            text="Map each shapefile bucket to its document section path (Heading 1 - Bold bullets)."
        ).pack(side='left', padx=5)

        self.mapping_button = ttk.Button(
            frame,
            text="Configure Mappings",
            command=self.open_mapping_dialog,
            state='normal'  # Always enabled now
        )
        self.mapping_button.pack(side='right', padx=5)

        # Status label
        self.mapping_status_var = tk.StringVar(value="No mappings configured")
        ttk.Label(frame, textvariable=self.mapping_status_var, foreground='gray').pack(
            side='right', padx=10
        )

    def setup_run_section(self, parent):
        """Setup run button section."""
        frame = ttk.Frame(parent, padding=10)
        frame.pack(fill='x', pady=10)

        self.run_button = ttk.Button(
            frame,
            text="4. Run Processing",
            command=self.run_processing,
            state='disabled'
        )
        self.run_button.pack()

        # Progress bar
        self.progress = ttk.Progressbar(frame, mode='indeterminate')
        self.progress.pack(fill='x', pady=5)

    def setup_output_section(self, parent):
        """Setup output/results section."""
        # Results notebook with tabs
        notebook = ttk.Notebook(parent)
        notebook.pack(fill='both', expand=True, pady=5)

        # Log tab
        log_frame = ttk.Frame(notebook)
        notebook.add(log_frame, text="Processing Log")

        self.log_text = scrolledtext.ScrolledText(
            log_frame,
            height=15,
            wrap=tk.WORD,
            font=('Courier', 9)
        )
        self.log_text.pack(fill='both', expand=True, padx=5, pady=5)

        # Unmatched settlements tab
        unmatched_frame = ttk.Frame(notebook)
        notebook.add(unmatched_frame, text="Unmatched Settlements")

        self.unmatched_text = scrolledtext.ScrolledText(
            unmatched_frame,
            height=15,
            wrap=tk.WORD,
            font=('Courier', 9)
        )
        self.unmatched_text.pack(fill='both', expand=True, padx=5, pady=5)

    def log(self, message):
        """Add message to log."""
        self.log_text.insert(tk.END, message + '\n')
        self.log_text.see(tk.END)
        self.root.update()

    def browse_document(self):
        """Browse for document file."""
        initial_dir = self.config.get('last_docx_dir', str(Path.home()))

        filename = filedialog.askopenfilename(
            title="Select Daily Collect Document",
            initialdir=initial_dir,
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )

        if filename:
            self.docx_path_var.set(filename)
            self.current_docx_path = filename
            self.config.set('last_docx_dir', str(Path(filename).parent))

            # Enable run button
            self.run_button.config(state='normal')

            self.log(f"Document selected: {os.path.basename(filename)}")

    def browse_output(self):
        """Browse for output directory."""
        initial_dir = self.output_path_var.get()
        if not initial_dir or not os.path.exists(initial_dir):
            # Default to document directory if available, otherwise home
            if self.current_docx_path:
                initial_dir = str(Path(self.current_docx_path).parent)
            else:
                initial_dir = str(Path.home())

        directory = filedialog.askdirectory(
            title="Select Output Directory",
            initialdir=initial_dir
        )

        if directory:
            self.output_path_var.set(directory)
            self.config.set('last_output_dir', directory)
            self.log(f"Output directory set to: {directory}")

    def browse_shapefile(self, key):
        """Browse for shapefile."""
        filename = filedialog.askopenfilename(
            title="Select Shapefile",
            filetypes=[("Shapefiles", "*.shp"), ("All Files", "*.*")]
        )

        if filename:
            self.shapefile_vars[key].set(filename)
            self.config.set(f'shapefile_paths.{key}', filename)
            self.log(f"Shapefile updated: {key} -> {os.path.basename(filename)}")

    def _filter_sections_by_range(self, doc, sections):
        """
        Filter sections to only include those between 'Axis-Specific Kinetic Collect'
        and 'Crimea' heading.

        Args:
            doc: Document object
            sections: List of section dictionaries from parse_document_structure

        Returns:
            List of filtered sections
        """
        # Find the paragraph indices for start and end markers
        start_index = None
        end_index = None

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip().lower()

            # Check for start marker: "Axis-Specific Kinetic Collect"
            if 'axis-specific kinetic collect' in text:
                start_index = i

            # Check for end marker: "Crimea" as Heading 1
            if para.style.name == 'Heading 1' and 'crimea' in text:
                end_index = i
                break

        # If we didn't find the markers, return all sections (fallback to original behavior)
        if start_index is None:
            return sections

        # Filter sections based on paragraph indices
        filtered = []
        for section in sections:
            # Get the paragraph index for this section
            para = section.get('paragraph')
            if para is not None:
                # Find the index of this paragraph in the document
                try:
                    para_index = doc.paragraphs.index(para)

                    # Include if after start marker
                    if start_index is not None and para_index <= start_index:
                        continue

                    # Exclude if at or after end marker
                    if end_index is not None and para_index >= end_index:
                        continue

                    filtered.append(section)
                except ValueError:
                    # Paragraph not found in document, skip it
                    continue

        return filtered

    def _extract_sitrep_bucket_paths(self, doc):
        """
        Bottom-up approach: Find UA GEN STAFF sitrep paragraphs, then work backwards
        to find all bold bullets above them to build the full nested path.

        Returns:
            set: Unique bucket paths (e.g., "Western Zaporizhia Oblast - General")
        """
        bucket_paths = set()
        current_heading1 = None
        sitreps_found = 0

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()

            if not text:
                continue

            # Track Heading 1 sections
            if para.style.name == 'Heading 1':
                current_heading1 = text
                continue

            # Found a UA GEN STAFF sitrep paragraph
            if 'ua gen staff sitrep' in text.lower():
                sitreps_found += 1
                # Work backwards from this paragraph to find all bold bullets above it
                bold_bullets = []

                # Search backwards through previous paragraphs (limit to avoid collecting too many)
                # We'll stop after collecting a reasonable number of bullets (max 3-4 levels of nesting)
                MAX_BULLETS_TO_COLLECT = 4
                bullets_collected = 0

                for j in range(i - 1, -1, -1):
                    prev_para = doc.paragraphs[j]
                    prev_text = prev_para.text.strip()

                    if not prev_text:
                        continue

                    # Stop if we hit a Heading 1 (we've gone too far back)
                    if prev_para.style.name == 'Heading 1':
                        break

                    # Check if this paragraph is a bold bullet
                    # A paragraph is a bold bullet if:
                    # 1. It has at least one bold run with actual text
                    # 2. It's formatted as a list OR starts with a bullet character
                    has_bold_text = False
                    for run in prev_para.runs:
                        if run.bold and run.text.strip():
                            has_bold_text = True
                            break

                    is_list_style = prev_para.style.name.startswith('List')
                    starts_with_bullet = prev_text.startswith(('•', '-', '·'))

                    is_bullet = is_list_style or starts_with_bullet

                    if has_bold_text and is_bullet:
                        # Clean bullet characters and extra whitespace
                        clean_text = re.sub(r'^[•\-·]\s*', '', prev_text).strip()
                        # Remove any trailing parenthetical notes
                        clean_text = re.sub(r'\s*\([^)]*\)\s*$', '', clean_text).strip()

                        # Only add if not empty after cleaning
                        if clean_text:
                            bold_bullets.insert(0, clean_text)  # Insert at beginning to maintain order
                            bullets_collected += 1

                            # Stop after collecting enough bullets to avoid collecting siblings
                            if bullets_collected >= MAX_BULLETS_TO_COLLECT:
                                break

                    # Don't stop at bold non-bullet text (e.g., "Luhansk Oblast" dividers)
                    # Continue searching backwards until we hit a Heading 1 or collect enough bullets

                # Build the full path from heading1 and bold bullets
                if current_heading1 and bold_bullets:
                    # Join all bold bullets with " - "
                    bullet_path = " - ".join(bold_bullets)
                    full_path = f"{current_heading1} - {bullet_path}"
                    bucket_paths.add(full_path)
                    print(f"DEBUG: Found bucket path: {full_path}")
                elif current_heading1:
                    # If no bold bullets found, just use heading1
                    full_path = f"{current_heading1} - General"
                    bucket_paths.add(full_path)
                    print(f"DEBUG: Found bucket path (no bullets): {full_path}")

        print(f"DEBUG: Total sitreps found: {sitreps_found}")
        print(f"DEBUG: Unique bucket paths: {len(bucket_paths)}")

        return bucket_paths

    def open_mapping_dialog(self):
        """Open the shapefile bucket to document section mapping dialog."""
        try:
            # Get bucket names from shapefile
            bucket_shp = self.shapefile_vars['bucket_shp'].get()
            if not os.path.exists(bucket_shp):
                messagebox.showerror("Error", f"Bucket shapefile not found: {bucket_shp}")
                return

            self.log("Loading bucket names from shapefile...")
            gdf_buckets = gpd.read_file(bucket_shp)
            bucket_field = self.config.get('bucket_field', 'Name')

            if bucket_field not in gdf_buckets.columns:
                messagebox.showerror(
                    "Error",
                    f"Field '{bucket_field}' not found in bucket shapefile.\n"
                    f"Available fields: {', '.join(gdf_buckets.columns)}"
                )
                return

            bucket_names = gdf_buckets[bucket_field].unique().tolist()
            self.log(f"Found {len(bucket_names)} buckets in shapefile")

            # Get current mappings (bucket -> document path)
            current_mappings = self.config.get('bucket_to_document_map', {})

            # Show dialog (user enters document paths for each shapefile bucket)
            dialog = HeadingMapperDialog(self.root, bucket_names, current_mappings)
            result = dialog.show()

            if result is not None:
                self.config.set('bucket_to_document_map', result)
                count = len(result)
                self.mapping_status_var.set(f"{count} mapping(s) configured")
                self.log(f"Saved {count} bucket-to-document mappings")
                messagebox.showinfo("Success", f"Saved {count} mapping(s).")

        except Exception as e:
            messagebox.showerror("Error", f"Failed to open mapping dialog:\n{str(e)}")

    def run_processing(self):
        """Run the fighting circles processing."""
        if self.processing:
            return

        if not self.current_docx_path:
            messagebox.showerror("Error", "Please select a document first.")
            return

        # Run in separate thread to prevent GUI freezing
        thread = threading.Thread(target=self.process_document)
        thread.daemon = True
        thread.start()

    def process_document(self):
        """Process the document (runs in separate thread)."""
        self.processing = True
        self.run_button.config(state='disabled')
        self.progress.start()

        try:
            # Clear previous output
            self.log_text.delete(1.0, tk.END)
            self.unmatched_text.delete(1.0, tk.END)

            self.log("=" * 80)
            self.log("FIGHTING CIRCLES GENERATOR")
            self.log("=" * 80)
            self.log(f"Document: {os.path.basename(self.current_docx_path)}")
            self.log("")

            # Get shapefile paths
            city_shp = self.shapefile_vars['city_shp'].get()
            frontline_shp = self.shapefile_vars['frontline_shp'].get()
            bucket_shp = self.shapefile_vars['bucket_shp'].get()

            # Verify files exist
            for path, name in [(city_shp, "City points"),
                              (frontline_shp, "Frontline"),
                              (bucket_shp, "Collection buckets")]:
                if not os.path.exists(path):
                    raise FileNotFoundError(f"{name} shapefile not found: {path}")

            bucket_field = self.config.get('bucket_field', 'Name')

            # Load document
            doc = Document(self.current_docx_path)

            # Step 1: Extract settlements
            self.log("[1/4] Extracting settlements from UA GEN STAFF sitreps...")
            settlements_data, special_settlements = extract_gen_staff_settlements(doc)

            if not settlements_data and not special_settlements:
                self.log("ERROR: No settlements found!")
                self.log("Make sure document contains 'UA GEN STAFF sitrep' text")
                return

            self.log(f"Found {len(settlements_data)} regular settlement mentions")
            if special_settlements:
                self.log(f"Found {len(special_settlements)} special settlements:")
                for special in special_settlements:
                    self.log(f"  • {special['name']}")

            # Step 1.5: Map document paths to shapefile buckets using user's manual mappings
            self.log("\n[1.5/4] Applying bucket mappings...")
            bucket_to_doc_map = self.config.get('bucket_to_document_map', {})

            if not bucket_to_doc_map:
                self.log("WARNING: No bucket mappings configured!")
                self.log("Please configure mappings before running.")
                messagebox.showwarning(
                    "Warning",
                    "No bucket mappings configured.\n\n"
                    "Please click 'Configure Mappings' to set up bucket-to-document mappings."
                )
                return

            # Create reverse mapping: document_path → bucket_name
            doc_to_bucket_map = {doc_path.strip(): bucket_name.strip()
                                for bucket_name, doc_path in bucket_to_doc_map.items()}

            self.log(f"Using {len(doc_to_bucket_map)} configured mappings")

            # Apply mappings to settlements
            mapped_count = 0
            unmapped_buckets = set()
            filtered_settlements = []

            for item in settlements_data:
                doc_bucket = item['bucket']  # e.g., "Kharkiv Oblast - Kupyansk direction"

                # Look up the shapefile bucket name
                if doc_bucket in doc_to_bucket_map:
                    item['bucket'] = doc_to_bucket_map[doc_bucket]  # Replace with shapefile bucket name
                    filtered_settlements.append(item)
                    mapped_count += 1
                else:
                    unmapped_buckets.add(doc_bucket)

            # Replace settlements_data with filtered version (only mapped settlements)
            settlements_data = filtered_settlements

            self.log(f"Mapped {mapped_count} settlements to shapefile buckets")

            if unmapped_buckets:
                self.log(f"WARNING: {len(unmapped_buckets)} document sections have no bucket mapping:")
                for bucket in sorted(unmapped_buckets):
                    self.log(f"  • {bucket}")
                self.log("These settlements will be skipped.")

            # Show sample
            if settlements_data:
                self.log("\nSample mapped settlements:")
                for item in settlements_data[:8]:
                    self.log(f"  • {item['settlement']:25} → {item['bucket']}")
                if len(settlements_data) > 8:
                    self.log(f"  ... and {len(settlements_data) - 8} more")
            else:
                self.log("\nERROR: No settlements were successfully mapped!")
                self.log("Check your bucket mappings and try again.")
                return

            # Step 2: Match to buckets
            self.log(f"\n[2/4] Matching settlements to collection buckets...")
            matched, unmatched = match_settlements_to_buckets(
                settlements_data, city_shp, frontline_shp, bucket_shp,
                bucket_name_field=bucket_field
            )

            # Add special settlements
            metric_crs = gpd.read_file(city_shp).estimate_utm_crs()
            from shapely.geometry import Point

            for special in special_settlements:
                point_wgs84 = Point(special['longitude'], special['latitude'])
                point_metric = gpd.GeoSeries([point_wgs84], crs='EPSG:4326').to_crs(metric_crs).iloc[0]

                matched.append({
                    'settlement': special['name'].lower(),
                    'matched_name': special['name'],
                    'geometry': point_metric,
                    'direction': 'Special',
                    'heading1': 'Special',
                    'count': 1,
                    'dist_to_front': 0,
                    'is_special': True
                })

            if not matched:
                self.log("ERROR: No settlements matched!")
                return

            special_count = len([m for m in matched if m.get('is_special', False)])
            regular_count = len(matched) - special_count

            self.log(f"Matched {len(matched)} settlements:")
            self.log(f"  - Regular: {regular_count}")
            self.log(f"  - Special: {special_count}")
            self.log(f"  - Unmatched: {len(unmatched)}")

            # Step 3: Generate fighting circles
            self.log(f"\n[3/4] Generating fighting circles...")
            gdf_circles = create_fighting_circles(matched, city_shp)

            multi_mention = len([m for m in matched if m['count'] > 1])
            if multi_mention:
                self.log(f"{multi_mention} settlements with multiple mentions (dispersed)")

            self.log(f"Generated {len(gdf_circles)} total fighting circles")

            # Step 4: Export
            self.log(f"\n[4/4] Exporting results...")
            today = datetime.now().strftime("%m%d%Y")

            # Use custom output directory if specified, otherwise use document directory
            custom_output = self.output_path_var.get()
            if custom_output and os.path.exists(custom_output):
                output_dir = Path(custom_output)
            else:
                output_dir = Path(self.current_docx_path).parent

            output_shp = output_dir / f"FC_{today}.shp"
            output_csv = output_dir / f"FC_{today}.csv"

            gdf_circles.to_file(str(output_shp))

            csv_data = gdf_circles.copy()
            csv_data = csv_data.drop(columns=['geometry'])
            csv_data.to_csv(str(output_csv), index=False)

            self.log(f"Shapefile: {output_shp.name}")
            self.log(f"CSV: {output_csv.name}")

            # Summary
            self.log(f"\n{'=' * 80}")
            self.log("SUMMARY")
            self.log(f"{'=' * 80}")
            self.log(f"Total circles: {len(gdf_circles)}")
            self.log(f"Unique settlements: {gdf_circles['settlement'].nunique()}")
            self.log(f"Axes/Oblasts: {gdf_circles['heading1'].nunique()}")

            self.log(f"\nBy Heading:")
            heading_counts = gdf_circles.groupby('heading1').size().sort_values(ascending=False)
            for heading, count in heading_counts.items():
                self.log(f"  {heading}: {count}")

            # Display unmatched settlements
            if unmatched:
                self.log(f"\nWARNING: {len(unmatched)} settlements unmatched (see Unmatched tab)")

                self.unmatched_text.insert(tk.END, f"Unmatched Settlements ({len(unmatched)} total)\n")
                self.unmatched_text.insert(tk.END, "=" * 80 + "\n\n")

                for item in unmatched:
                    self.unmatched_text.insert(
                        tk.END,
                        f"Settlement: {item['settlement']}\n"
                        f"Direction: {item['direction']}\n"
                        f"Heading: {item['heading1']}\n"
                        f"Source: {item['source_text'][:100]}...\n"
                        f"{'-' * 80}\n"
                    )
            else:
                self.unmatched_text.insert(tk.END, "No unmatched settlements! All settlements were successfully matched.")

            self.log(f"\n{'=' * 80}")
            self.log("PROCESSING COMPLETE")
            self.log(f"{'=' * 80}")

            messagebox.showinfo(
                "Success",
                f"Processing complete!\n\n"
                f"Generated {len(gdf_circles)} fighting circles\n"
                f"Unmatched: {len(unmatched)}\n\n"
                f"Files saved to:\n{output_dir}"
            )

        except Exception as e:
            self.log(f"\nERROR: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
            messagebox.showerror("Error", f"Processing failed:\n{str(e)}")

        finally:
            self.processing = False
            self.run_button.config(state='normal')
            self.progress.stop()


def main():
    """Main entry point."""
    root = tk.Tk()
    app = FightingCirclesGUI(root)
    root.mainloop()


if __name__ == "__main__":
    main()
