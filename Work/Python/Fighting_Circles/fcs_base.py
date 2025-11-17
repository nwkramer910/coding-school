from docx import Document
import pandas as pd
import geopandas as gpd
import os
from rapidfuzz import process, fuzz
from shapely.geometry import Point
import numpy as np
from datetime import datetime
import re
from collections import defaultdict



def parse_document_structure(doc):
    """
    Parse document structure using:
    - Heading 1 for major axes/oblasts
    - Bullet depth for sub-directions
    - Track hierarchy for bucket assignment
    """
    sections = []
    current_heading1 = None
    current_direction = None
    hierarchy = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if not text:
            continue
        
        # Check if Heading 1 (major axis/oblast)
        if para.style.name == 'Heading 1':
            current_heading1 = text
            current_direction = "General"
            hierarchy = [current_heading1]
            continue
        
        # Check if this is a bulleted/direction line
        if para.style.name.startswith('List') or text.startswith(('•', '-', '·')):
            # Clean bullet characters
            clean_text = re.sub(r'^[•\-·]\s*', '', text).strip()
            
            # Check if bold (indicates a section header)
            if any(run.bold for run in para.runs if run.text.strip()):
                current_direction = clean_text
                continue
        
        # This is content - check if it contains a sitrep
        if current_heading1:
            sections.append({
                'paragraph': para,
                'heading1': current_heading1,
                'direction': current_direction,
                'bucket': f"{current_heading1} - {current_direction}",
                'text': text
            })
    
    return sections


def extract_settlement_names(text):
    """
    Extract settlement names from UA GEN STAFF sitrep text.
    Handles patterns:
    - "IVO Settlement1, Settlement2, and Settlement3"
    - "toward Settlement4, Settlement5"
    - "IVO Settlement and toward Settlement2"
    """
    settlements = []
    
    # Remove URLs
    text = re.sub(r'https?://\S+', '', text)
    
    # Pattern 1: Extract everything between "IVO" and next delimiter
    # Delimiters: "toward", "and toward", period, end of string
    ivo_pattern = r'IVO\s+([\w\s,\-]+?)(?=\s+and\s+toward|,\s+and\s+toward|\s+toward|\.|$)'
    ivo_matches = re.finditer(ivo_pattern, text, re.IGNORECASE)
    
    for match in ivo_matches:
        settlement_string = match.group(1)
        # Split by commas and " and "
        parts = re.split(r',\s*|\s+and\s+', settlement_string)
        for part in parts:
            clean = part.strip().lower()
            # Filter: no digits, reasonable length, not common words
            if (clean and 
                len(clean) > 2 and 
                not re.search(r'\d', clean) and
                clean not in ['times', 'attacked', 'reported', 'c', 'the']):
                settlements.append(clean)
    
    # Pattern 2: Extract everything after "toward"
    toward_pattern = r'toward\s+([\w\s,\-]+?)(?=\s+and\s+IVO|,\s+and\s+IVO|\s+IVO|\.|$)'
    toward_matches = re.finditer(toward_pattern, text, re.IGNORECASE)
    
    for match in toward_matches:
        settlement_string = match.group(1)
        parts = re.split(r',\s*|\s+and\s+', settlement_string)
        for part in parts:
            clean = part.strip().lower()
            if (clean and 
                len(clean) > 2 and 
                not re.search(r'\d', clean) and
                clean not in ['times', 'attacked', 'reported', 'c', 'the']):
                settlements.append(clean)
    
    # Remove duplicates while preserving order
    seen = set()
    unique_settlements = []
    for s in settlements:
        if s not in seen:
            seen.add(s)
            unique_settlements.append(s)
    
    return unique_settlements

def extract_special_settlements(doc):
    """
    Extract special hardcoded settlements that don't follow normal IVO/toward patterns.
    These can appear anywhere in the document.
    
    Searches for specific phrases and returns their hardcoded coordinates.
    """
    special_found = []
    
    # Define search patterns and their coordinates (lat, lon)
    special_patterns = {
        r'north\s+slobozhansk': ('North Slobozhansk', 51.1207270, 34.9237297),
        r'kursk\s+dir(?:s|ection)?': ('Kursk', 51.1207270, 34.9237297),
        r'dnipro\s+dir(?:ection)?': ('Dnipro', 46.5262724, 32.3564850),
        r'(?:IVO\s+)?antonyvskyi\s+bridge': ('Antonyvskyi Bridge', 46.6709284, 32.7206244),
        r'(?:IVO\s+)?antonivskyi\s+bridge': ('Antonivskyi Bridge', 46.6709284, 32.7206244),
    }
    
    # Search through all paragraphs (no structure parsing needed)
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Check each special pattern
        for pattern, (name, lat, lon) in special_patterns.items():
            if re.search(pattern, text, re.IGNORECASE):
                special_found.append({
                    'name': name,
                    'latitude': lat,
                    'longitude': lon,
                    'source_text': text[:150]  # For debugging
                })
    
    return special_found

def extract_gen_staff_settlements(doc):
    """
    Extract settlements from paragraphs containing UA GEN STAFF sitrep text.
    Returns both regular settlements and special hardcoded settlements.
    """
    sections = parse_document_structure(doc)
    settlements_by_bucket = []
    
    for section in sections:
        para = section['paragraph']
        para_text = section['text']
        
        # Check if paragraph contains UA GEN STAFF sitrep
        if not re.search(r'UA\s+GEN\s+STAFF\s+sitrep', para_text, re.IGNORECASE):
            continue
        
        # Extract settlement names
        settlements = extract_settlement_names(para_text)
        
        for settlement in settlements:
            settlements_by_bucket.append({
                'settlement': settlement,
                'heading1': section['heading1'],
                'direction': section['direction'],
                'bucket': section['bucket'],
                'source_text': para_text[:150]
            })
    
    # Extract special settlements (independent of structure)
    special_settlements = extract_special_settlements(doc)
    
    return settlements_by_bucket, special_settlements


def normalize_bucket_name(direction, heading1):
    """
    Normalize direction names to match bucket shapefile names.
    
    Mapping logic:
    - "General" under "Kharkiv Oblast" → "Kharkiv Oblast"
    - "Kupyansk (international border to Hlushkivka)" → "Kupyansk"
    - "Lyman direction" → "Lyman"
    - etc.
    """
    # If General, use the heading1 (oblast/axis name)
    if direction == "General":
        return heading1
    
    # Remove parenthetical descriptions
    clean = re.sub(r'\s*\([^)]*\)', '', direction).strip()
    
    # Remove common suffixes
    clean = re.sub(r'\s+(direction|area|axis|salient|effort)$', '', clean, flags=re.IGNORECASE)
    
    return clean


def match_settlements_to_buckets(settlements_data, city_shp, frontline_shp, bucket_shp, 
                                   bucket_name_field='BUCKET', distance_threshold=15000):
    """
    Match settlements to bucket polygons with special settlement handling.
    """
    # Load spatial data
    gdf_cities = gpd.read_file(city_shp)
    gdf_frontline = gpd.read_file(frontline_shp)
    gdf_buckets = gpd.read_file(bucket_shp)
    
    # Standardize names
    gdf_cities['ADM4_EN'] = gdf_cities['ADM4_EN'].str.lower().str.strip()
    gdf_buckets[bucket_name_field] = gdf_buckets[bucket_name_field].str.strip()
    
    # Project to metric CRS
    metric_crs = gdf_cities.estimate_utm_crs()
    gdf_cities_proj = gdf_cities.to_crs(metric_crs)
    gdf_frontline_proj = gdf_frontline.to_crs(metric_crs)
    gdf_buckets_proj = gdf_buckets.to_crs(metric_crs)
    
    frontline_union = gdf_frontline_proj.geometry.unary_union
    
    matched_settlements = []
    unmatched = []
    special_matched = []
    
    # Count occurrences
    settlement_counts = defaultdict(int)
    for item in settlements_data:
        key = f"{item['settlement']}|{item['direction']}"
        settlement_counts[key] += 1
    
    for item in settlements_data:
        settlement_name = item['settlement']
        direction = item['direction']
        heading1 = item['heading1']
        
        # Normalize bucket name for matching
        search_name = normalize_bucket_name(direction, heading1)
        
        # Find matching bucket
        bucket_geom = None
        bucket_names = gdf_buckets_proj[bucket_name_field].tolist()
        
        # Try exact match first
        exact_match = gdf_buckets_proj[gdf_buckets_proj[bucket_name_field] == search_name]
        if not exact_match.empty:
            bucket_geom = exact_match
        else:
            # Fuzzy match
            match_result = process.extractOne(search_name, bucket_names, scorer=fuzz.ratio)
            if match_result and match_result[1] >= 60:
                matched_bucket_name = match_result[0]
                bucket_geom = gdf_buckets_proj[gdf_buckets_proj[bucket_name_field] == matched_bucket_name]
        
        if bucket_geom is None or bucket_geom.empty:
            unmatched.append(item)
            print(f"  ⚠ No bucket found for '{direction}' (searched: '{search_name}')")
            continue
        
        bucket_polygon = bucket_geom.geometry.unary_union
        
        # Find candidate cities within bucket
        candidates = gdf_cities_proj[gdf_cities_proj.within(bucket_polygon)]
        
        if candidates.empty:
            unmatched.append(item)
            print(f"  ⚠ No cities in bucket for '{settlement_name}' in {direction}")
            continue
        
        # Match settlement name
        candidate_names = candidates['ADM4_EN'].tolist()
        
        # Try exact match
        exact_settlement = candidates[candidates['ADM4_EN'] == settlement_name]
        
        if not exact_settlement.empty:
            matched_cities = exact_settlement.copy()
        else:
            # Fuzzy match
            match_result = process.extractOne(settlement_name, candidate_names, scorer=fuzz.ratio)
            
            if match_result and match_result[1] >= 60:
                matched_name = match_result[0]
                matched_cities = candidates[candidates['ADM4_EN'] == matched_name].copy()
            else:
                unmatched.append(item)
                print(f"  ⚠ No match for '{settlement_name}' in {direction}")
                continue
        
        # Calculate distance to frontline
        matched_cities['dist_to_front'] = matched_cities.geometry.apply(
            lambda geom: frontline_union.distance(geom)
        )
        
        # Filter by distance threshold
        within_threshold = matched_cities[matched_cities['dist_to_front'] <= distance_threshold]
        
        if not within_threshold.empty:
            # Get closest to frontline
            within_threshold = within_threshold.sort_values('dist_to_front')
            best_match = within_threshold.iloc[0]
            
            count_key = f"{settlement_name}|{direction}"
            count = settlement_counts[count_key]
            
            matched_settlements.append({
                'settlement': settlement_name,
                'matched_name': best_match['ADM4_EN'],
                'geometry': best_match.geometry,
                'direction': direction,
                'heading1': heading1,
                'count': count,
                'dist_to_front': best_match['dist_to_front'],
                'is_special': False
            })
        else:
            unmatched.append(item)
            print(f"  ⚠ '{settlement_name}' outside {distance_threshold/1000}km threshold")
    
    # Combine regular and special settlements
    all_matched = matched_settlements + special_matched
    
    return all_matched, unmatched


def create_fighting_circles(matched_settlements, city_shp):
    """
    Generate fighting circles with dispersion for multiple mentions.
    Special settlements use exact coordinates without dispersion.
    """
    gdf_cities = gpd.read_file(city_shp)
    gdf_cities['ADM4_EN'] = gdf_cities['ADM4_EN'].str.lower().str.strip()
    
    metric_crs = gdf_cities.estimate_utm_crs()
    gdf_cities_proj = gdf_cities.to_crs(metric_crs)
    
    fighting_circles = []
    
    for item in matched_settlements:
        settlement_name = item['matched_name']
        count = item['count']
        centroid = item['geometry']
        is_special = item.get('is_special', False)
        
        # Special settlements: no dispersion, just duplicate at exact coordinates
        if is_special:
            for i in range(count):
                fighting_circles.append({
                    'settlement': settlement_name,
                    'geometry': centroid,
                    'direction': item['direction'],
                    'heading1': item['heading1'],
                    'mention_num': i + 1,
                    'total_mentions': count,
                    'dist_to_front_m': item['dist_to_front'],
                    'is_special': True
                })
            continue
        
        # Regular settlements: disperse if multiple mentions
        if count == 1:
            fighting_circles.append({
                'settlement': settlement_name,
                'geometry': centroid,
                'direction': item['direction'],
                'heading1': item['heading1'],
                'mention_num': 1,
                'total_mentions': 1,
                'dist_to_front_m': item['dist_to_front'],
                'is_special': False
            })
        else:
            # Get settlement boundary for dispersion
            settlement_geom = gdf_cities_proj[gdf_cities_proj['ADM4_EN'] == item['settlement']]
            
            if not settlement_geom.empty:
                settlement_poly = settlement_geom.geometry.iloc[0]
                
                if hasattr(settlement_poly, 'bounds'):
                    bounds = settlement_poly.bounds
                    width = bounds[2] - bounds[0]
                    height = bounds[3] - bounds[1]
                    radius = min(width, height) * 0.3  # 30% of smaller dimension
                    
                    for i in range(count):
                        angle = (2 * np.pi * i) / count
                        offset_x = centroid.x + radius * np.cos(angle)
                        offset_y = centroid.y + radius * np.sin(angle)
                        dispersed_point = Point(offset_x, offset_y)
                        
                        # Use dispersed point if within bounds, else centroid
                        if hasattr(settlement_poly, 'contains') and settlement_poly.contains(dispersed_point):
                            use_point = dispersed_point
                        else:
                            use_point = centroid
                        
                        fighting_circles.append({
                            'settlement': settlement_name,
                            'geometry': use_point,
                            'direction': item['direction'],
                            'heading1': item['heading1'],
                            'mention_num': i + 1,
                            'total_mentions': count,
                            'dist_to_front_m': item['dist_to_front'],
                            'is_special': False
                        })
                else:
                    # Point geometry - duplicate at centroid
                    for i in range(count):
                        fighting_circles.append({
                            'settlement': settlement_name,
                            'geometry': centroid,
                            'direction': item['direction'],
                            'heading1': item['heading1'],
                            'mention_num': i + 1,
                            'total_mentions': count,
                            'dist_to_front_m': item['dist_to_front'],
                            'is_special': False
                        })
    
    # Convert to geographic CRS
    gdf_circles_proj = gpd.GeoDataFrame(fighting_circles, geometry='geometry', crs=metric_crs)
    gdf_circles = gdf_circles_proj.to_crs(gdf_cities.crs)
    
    # Add lat/lon
    gdf_circles['latitude'] = gdf_circles.geometry.y
    gdf_circles['longitude'] = gdf_circles.geometry.x
    
    return gdf_circles


def main():
    """Main execution function."""
    print("=" * 80)
    print("FIGHTING CIRCLES GENERATOR - UA GEN STAFF Sitrep Edition")
    print("=" * 80)
    
    # Get document path
    while True:
        docx_path = input("\nEnter path to daily collect document (.docx): ").strip().strip('"').strip("'")
        if os.path.exists(docx_path) and docx_path.lower().endswith('.docx'):
            break
        print("  ✗ Error: Invalid file path. Please try again.")
    
    # Shapefile paths
    city_shp = "citypoints.shp"
    frontline_shp = "frontline_frontline.shp"
    bucket_shp = "Collection Buckets.shp"
    
    # Verify shapefiles
    for shp_var, shp_name in [('city_shp', "city points"), 
                               ('frontline_shp', "frontline"),
                               ('bucket_shp', "collection buckets")]:
        shp_path = locals()[shp_var]
        if not os.path.exists(shp_path):
            new_path = input(f"  Enter path to {shp_name} shapefile: ").strip().strip('"').strip("'")
            if shp_name == "city points":
                city_shp = new_path
            elif shp_name == "frontline":
                frontline_shp = new_path
            else:
                bucket_shp = new_path
    
    # Get bucket field name
    bucket_field = 'Name'
    
    print(f"\n{'=' * 80}")
    print(f"Processing: {os.path.basename(docx_path)}")
    print(f"{'=' * 80}\n")
    
    try:
        # Load document
        doc = Document(docx_path)
        
        # Step 1: Extract settlements
        print("[1/4] Extracting settlements from UA GEN STAFF sitreps...")
        settlements_data, special_settlements = extract_gen_staff_settlements(doc)

        if not settlements_data and not special_settlements:
            print("\n  ✗ ERROR: No settlements found!")
            print("    Make sure document contains 'UA GEN STAFF sitrep' text or special settlement phrases")
            return

        print(f"  ✓ Found {len(settlements_data)} regular settlement mentions")
        if special_settlements:
            print(f"  ✓ Found {len(special_settlements)} special settlements:")
            for special in special_settlements:
                print(f"    • {special['name']}")

        # Show sample regular settlements
        if settlements_data:
            print("\n  Sample extractions:")
            for item in settlements_data[:8]:
                print(f"    • {item['settlement']:25} → {item['bucket']}")
            if len(settlements_data) > 8:
                print(f"    ... and {len(settlements_data) - 8} more")
        
        # Step 2: Match to buckets
        print(f"\n[2/4] Matching settlements to collection buckets...")
        matched, unmatched = match_settlements_to_buckets(
            settlements_data, city_shp, frontline_shp, bucket_shp,
            bucket_name_field=bucket_field
        )

        # Step 2.5: Add special settlements directly (no bucket matching needed)
        metric_crs = gpd.read_file(city_shp).estimate_utm_crs()
        for special in special_settlements:
            point_wgs84 = Point(special['longitude'], special['latitude'])
            point_metric = gpd.GeoSeries([point_wgs84], crs='EPSG:4326').to_crs(metric_crs).iloc[0]
            
            matched.append({
                'settlement': special['name'].lower(),
                'matched_name': special['name'],
                'geometry': point_metric,
                'direction': 'Special',
                'heading1': 'Special',
                'count': 1,  # Always count as 1 for special settlements
                'dist_to_front': 0,
                'is_special': True
            })

        if not matched:
            print("\n  ✗ ERROR: No settlements matched!")
            print(f"    Check bucket shapefile field name: '{bucket_field}'")
            return

        special_count = len([m for m in matched if m.get('is_special', False)])
        regular_count = len(matched) - special_count

        print(f"  ✓ Matched {len(matched)} settlements:")
        print(f"    - Regular: {regular_count}")
        print(f"    - Special: {special_count}")
        print(f"    - Unmatched: {len(unmatched)}")
        
        # Step 3: Generate fighting circles
        print(f"\n[3/4] Generating fighting circles...")
        gdf_circles = create_fighting_circles(matched, city_shp)
        
        multi_mention = len([m for m in matched if m['count'] > 1])
        if multi_mention:
            print(f"  ✓ {multi_mention} settlements with multiple mentions (dispersed)")
        
        print(f"  ✓ Generated {len(gdf_circles)} total fighting circles")
        
        # Step 4: Export
        print(f"\n[4/4] Exporting results...")
        today = datetime.now().strftime("%m%d%Y")
        output_shp = f"FC_{today}.shp"
        output_csv = f"FC_{today}.csv"
        
        gdf_circles.to_file(output_shp)
        
        # CSV export
        csv_data = gdf_circles.copy()
        csv_data = csv_data.drop(columns=['geometry'])
        csv_data.to_csv(output_csv, index=False)
        
        print(f"  ✓ Shapefile: {output_shp}")
        print(f"  ✓ CSV: {output_csv}")
        
        # Summary
        print(f"\n{'=' * 80}")
        print("SUMMARY")
        print(f"{'=' * 80}")
        print(f"Total circles: {len(gdf_circles)}")
        print(f"Unique settlements: {gdf_circles['settlement'].nunique()}")
        print(f"Axes/Oblasts: {gdf_circles['heading1'].nunique()}")
        
        print(f"\nBy Heading:")
        heading_counts = gdf_circles.groupby('heading1').size().sort_values(ascending=False)
        for heading, count in heading_counts.items():
            print(f"  {heading}: {count}")
        
        # Unmatched warning
        if unmatched:
            print(f"\n⚠  WARNING: {len(unmatched)} settlements unmatched:")
            for item in unmatched[:15]:
                print(f"  • {item['settlement']} (in {item['direction']})")
            if len(unmatched) > 15:
                print(f"  ... and {len(unmatched) - 15} more")
        
        print(f"\n{'=' * 80}")
        print("✓ PROCESSING COMPLETE")
        print(f"{'=' * 80}\n")
        
    except Exception as e:
        print(f"\n✗ ERROR: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()